import os
from datetime import datetime

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUTPUT_DIR = "output"


def generate_document(user_request: str, sections: list) -> str:
    if not os.path.exists(OUTPUT_DIR):
        os.makedirs(OUTPUT_DIR)

    document = Document()

    _add_title(document, user_request)

    date_line = document.add_paragraph(
        f"Generated on {datetime.now().strftime('%d %B %Y, %H:%M')}"
    )
    date_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_line.runs[0].italic = True

    document.add_paragraph("")

    for section in sections:
        document.add_heading(section["title"], level=1)
        document.add_paragraph(section["content"])

    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    file_path = os.path.join(OUTPUT_DIR, f"document_{timestamp}.docx")

    document.save(file_path)
    return file_path


def _add_title(document: Document, user_request: str) -> None:
    short_request = user_request.strip()
    if len(short_request) > 80:
        short_request = short_request[:77] + "..."

    title_paragraph = document.add_paragraph()
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run = title_paragraph.add_run(short_request)
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(0x0B, 0x1A, 0x3A)